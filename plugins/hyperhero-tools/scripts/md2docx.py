#!/usr/bin/env python3
"""Convert a Markdown user guide to DOCX with embedded images and CJK support.

Usage:
    python md2docx.py <input.md> <output.docx> <screenshots-dir>
"""
import re
import os
import sys
import platform

TICK = chr(96)
TRIPLE = TICK * 3
INLINE_PAT = re.compile(
    r"(\*\*.*?\*\*|\*[^*]+?\*|" + TICK + r"[^" + TICK + r"]+" + TICK + r"|\[.*?\]\(.*?\))"
)
LIST_RULES = [
    (re.compile(r"^\s+(\d+)\.\s+(.*)"), "List Number 2", 2),
    (re.compile(r"^(\d+)\.\s+(.*)"), "List Number", 2),
    (re.compile(r"^\s+[-*]\s+(.*)"), "List Bullet 2", 1),
    (re.compile(r"^[-*]\s+(.*)"), "List Bullet", 1),
]


def get_cjk_font():
    s = platform.system()
    return {"Windows": "Microsoft JhengHei", "Darwin": "PingFang TC"}.get(s, "Noto Sans CJK TC")


def ensure_docx():
    try:
        from docx import Document  # noqa: F401
        return True
    except ImportError:
        print("ERROR: python-docx required. Install: pip install python-docx")
        sys.exit(1)


class DocxConverter:
    def __init__(self, img_dir):
        from docx import Document
        from docx.shared import Cm, Pt, RGBColor
        from docx.oxml.ns import qn

        self.img_dir = img_dir
        self.doc = Document()
        cjk = get_cjk_font()

        for sec in self.doc.sections:
            sec.top_margin = sec.bottom_margin = Cm(2)
            sec.left_margin = sec.right_margin = Cm(2.5)

        ns = self.doc.styles["Normal"]
        ns.font.name = cjk
        ns.font.size = Pt(11)
        ns.element.rPr.rFonts.set(qn("w:eastAsia"), cjk)
        for lv in range(1, 5):
            hs = self.doc.styles["Heading %d" % lv]
            hs.font.name = cjk
            hs.element.rPr.rFonts.set(qn("w:eastAsia"), cjk)

        self.colors = {
            "red": RGBColor(0xFF, 0, 0),
            "gray": RGBColor(0x55, 0x55, 0x55),
            "link": RGBColor(0, 0x66, 0xCC),
            "code": RGBColor(0x88, 0, 0),
        }

    def apply_inline(self, para, text):
        from docx.shared import Pt
        for part in INLINE_PAT.split(text):
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                para.add_run(part[2:-2]).bold = True
            elif part.startswith("*") and part.endswith("*"):
                para.add_run(part[1:-1]).italic = True
            elif part[0] == TICK and part[-1] == TICK:
                r = para.add_run(part[1:-1])
                r.font.name = "Consolas"
                r.font.size = Pt(10)
                r.font.color.rgb = self.colors["code"]
            elif part.startswith("[") and "](" in part:
                m = re.match(r"\[(.*?)\]\((.*?)\)", part)
                if m:
                    r = para.add_run(m.group(1))
                    r.font.color.rgb = self.colors["link"]
                    r.underline = True
                else:
                    para.add_run(part)
            else:
                para.add_run(part)

    def add_image(self, alt, filename):
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        fp = os.path.join(self.img_dir, filename)
        if os.path.exists(fp):
            try:
                self.doc.add_picture(fp, width=Inches(5.5))
                self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                return
            except Exception:
                pass
        r = self.doc.add_paragraph().add_run("[" + alt + " - missing]")
        r.italic = True
        r.font.color.rgb = self.colors["red"]

    def flush_blockquote(self, bq_lines):
        from docx.shared import Cm
        if not bq_lines:
            return
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        self.apply_inline(p, " ".join(bq_lines))
        for r in p.runs:
            r.font.color.rgb = self.colors["gray"]

    def flush_table(self, rows):
        from docx.shared import Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT
        if len(rows) < 3:
            return
        headers, data = rows[0], rows[2:]
        tbl = self.doc.add_table(rows=1 + len(data), cols=len(headers))
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, val in enumerate(headers):
            tbl.rows[0].cells[j].text = ""
            r = tbl.rows[0].cells[j].paragraphs[0].add_run(val.strip())
            r.bold = True
            r.font.size = Pt(10)
        for ri, row in enumerate(data):
            for j, val in enumerate(row):
                if j >= len(headers):
                    continue
                cell = tbl.rows[ri + 1].cells[j]
                cell.text = ""
                self.apply_inline(cell.paragraphs[0], val.strip())
                for r in cell.paragraphs[0].runs:
                    if r.font.size is None:
                        r.font.size = Pt(10)
        self.doc.add_paragraph("")

    def add_code_block(self, code_lines):
        from docx.shared import Cm, Pt
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        r = p.add_run("\n".join(code_lines))
        r.font.name = "Consolas"
        r.font.size = Pt(9)

    def process_line(self, ln, state):
        if ln.startswith(TRIPLE):
            if state["in_code"]:
                self.add_code_block(state["code_buf"])
                state["code_buf"] = []
            state["in_code"] = not state["in_code"]
            return True

        if state["in_code"]:
            state["code_buf"].append(ln)
            return True

        if state["in_bq"] and not ln.startswith(">"):
            self.flush_blockquote(state["bq_buf"])
            state["bq_buf"] = []
            state["in_bq"] = False

        if "|" in ln and ln.strip().startswith("|"):
            cells = [c.strip() for c in ln.strip().strip("|").split("|")]
            state["tbl_rows"].append(cells)
            state["in_table"] = True
            return True

        if state["in_table"]:
            self.flush_table(state["tbl_rows"])
            state["tbl_rows"] = []
            state["in_table"] = False
            return False  # re-process current line

        if not ln.strip() or ln.strip() == "---":
            return True

        m = re.match(r"^(#{1,4})\s+(.*)", ln)
        if m:
            self.doc.add_heading(m.group(2), level=len(m.group(1)))
            return True

        m = re.match(r"^!\[(.*?)\]\((.*?)\)", ln)
        if m:
            self.add_image(m.group(1), os.path.basename(m.group(2)))
            return True

        if ln.startswith(">"):
            state["in_bq"] = True
            state["bq_buf"].append(ln.lstrip(">").strip())
            return True

        for rule_re, style, grp in LIST_RULES:
            m = rule_re.match(ln)
            if m:
                self.apply_inline(self.doc.add_paragraph(style=style), m.group(grp))
                return True

        self.apply_inline(self.doc.add_paragraph(), ln)
        return True

    def convert(self, lines, out_path):
        state = {"tbl_rows": [], "in_table": False, "in_code": False,
                 "code_buf": [], "in_bq": False, "bq_buf": []}
        i = 0
        while i < len(lines):
            ln = lines[i].rstrip("\n")
            advanced = self.process_line(ln, state)
            if advanced:
                i += 1

        if state["in_bq"]:
            self.flush_blockquote(state["bq_buf"])
        if state["in_table"]:
            self.flush_table(state["tbl_rows"])

        self.doc.save(out_path)
        return os.path.getsize(out_path)


def main():
    ensure_docx()
    if len(sys.argv) < 4:
        print(__doc__)
        sys.exit(1)

    md_path, out_path, img_dir = sys.argv[1], sys.argv[2], sys.argv[3]
    if not os.path.exists(md_path):
        print(f"ERROR: Input file not found: {md_path}")
        sys.exit(1)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    converter = DocxConverter(img_dir)
    size = converter.convert(lines, out_path)
    print(f"DOCX saved: {out_path} ({size:,} bytes)")


if __name__ == "__main__":
    main()
